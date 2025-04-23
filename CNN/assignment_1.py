from docx import Document

# Create a new Word document
doc = Document()

# Title
doc.add_heading('Perceptron Learning Algorithm Solution', level=1)

# Section: Given Data
doc.add_heading('Given Data:', level=2)
doc.add_paragraph("Initial Weights:")
doc.add_paragraph("  W1 = 0.1")
doc.add_paragraph("  W2 = 0.2")
doc.add_paragraph("  Wb = -0.2 (bias weight)")
doc.add_paragraph("Threshold = 0.2")
doc.add_paragraph("Learning Rate (α) = 0.1")
doc.add_paragraph("Activation Function: Step Function")
doc.add_paragraph("  y = 1 if net ≥ threshold, otherwise y = 0")

# Section: Step 1 - Compute Net Input and Output (y)
doc.add_heading('Step 1: Compute Net Input and Output (y)', level=2)
doc.add_paragraph("The net input (net) is calculated as:")
doc.add_paragraph("  net = (x1 * W1) + (x2 * W2) + (1 * Wb)")
doc.add_paragraph("")

# Table: Initial Computations
doc.add_paragraph("Initial Computations:")
table = doc.add_table(rows=5, cols=11)
table.style = 'Table Grid'
headers = ["x1", "x2", "Bias", "W1", "W2", "Wb", "Net Calculation", "Net", "y", "Target (t)", "Error (e = t - y)"]
for i, header in enumerate(headers):
    table.cell(0, i).text = header

data = [
    [0, 0, 1, 0.1, 0.2, -0.2, "-0.2", -0.2, "0", "0", "0"],
    [0, 1, 1, 0.1, 0.2, -0.2, "0.0", 0.0, "0", "0", "0"],
    [1, 0, 1, 0.1, 0.2, -0.2, "-0.1", -0.1, "0", "0", "0"],
    [1, 1, 1, 0.1, 0.2, -0.2, "0.1", 0.1, "0", "1", "1"]
]

for row_idx, row in enumerate(data, start=1):
    cells = table.rows[row_idx].cells
    for col_idx, value in enumerate(row):
        cells[col_idx].text = str(value)

# Section: Step 2 - Weight Update Using Perceptron Rule
doc.add_heading('Step 2: Weight Update Using Perceptron Rule', level=2)
doc.add_paragraph("For input (1,1), the output y = 0 (since net = 0.1 < threshold 0.2) but the target t = 1.")
doc.add_paragraph("Thus, Error (e) = 1 - 0 = 1.")
doc.add_paragraph("Weight updates for (1,1):")
doc.add_paragraph("  W1 = 0.1 + (0.1 * 1 * 1) = 0.2")
doc.add_paragraph("  W2 = 0.2 + (0.1 * 1 * 1) = 0.3")
doc.add_paragraph("  Wb = -0.2 + (0.1 * 1) = -0.1")

# Section: Step 3 - Recalculate with Updated Weights
doc.add_heading('Step 3: Recalculate with Updated Weights', level=2)
doc.add_paragraph("Using updated weights: W1 = 0.2, W2 = 0.3, Wb = -0.1")
table2 = doc.add_table(rows=5, cols=9)
table2.style = 'Table Grid'
headers2 = ["x1", "x2", "Bias", "Net Calculation", "Net", "y", "Target (t)", "Error (e)", "Comments"]
for i, header in enumerate(headers2):
    table2.cell(0, i).text = header

data2 = [
    [0, 0, 1, "-0.1", -0.1, "0", "0", "0", ""],
    [0, 1, 1, "0.2", 0.2, "1", "0", "-1", "Misclassified"],
    [1, 0, 1, "0.1", 0.1, "0", "0", "0", ""],
    [1, 1, 1, "0.4", 0.4, "1", "1", "0", ""]
]

for row_idx, row in enumerate(data2, start=1):
    cells = table2.rows[row_idx].cells
    for col_idx, value in enumerate(row):
        cells[col_idx].text = str(value)

# Section: Step 4 - Further Updates for Misclassified (0,1)
doc.add_heading('Step 4: Further Updates for Misclassified (0,1)', level=2)
doc.add_paragraph("For input (0,1), error e = 0 - 1 = -1:")
doc.add_paragraph("  W1 = 0.2 + (0.1 * 0 * -1) = 0.2")
doc.add_paragraph("  W2 = 0.3 + (0.1 * 1 * -1) = 0.2")
doc.add_paragraph("  Wb = -0.1 + (0.1 * -1) = -0.2")

# Section: Step 5 - Final Check
doc.add_heading('Step 5: Final Check', level=2)
doc.add_paragraph("Final Weights: W1 = 0.2, W2 = 0.2, Wb = -0.2")
table3 = doc.add_table(rows=5, cols=9)
table3.style = 'Table Grid'
headers3 = ["x1", "x2", "Bias", "Net Calculation", "Net", "y", "Target (t)", "Error", "Comments"]
for i, header in enumerate(headers3):
    table3.cell(0, i).text = header

data3 = [
    [0, 0, 1, "-0.2", -0.2, "0", "0", "0", ""],
    [0, 1, 1, "0.0", 0.0, "0", "0", "0", ""],
    [1, 0, 1, "0.0", 0.0, "0", "0", "0", ""],
    [1, 1, 1, "0.2", 0.2, "1", "1", "0", ""]
]

for row_idx, row in enumerate(data3, start=1):
    cells = table3.rows[row_idx].cells
    for col_idx, value in enumerate(row):
        cells[col_idx].text = str(value)

# Section: Final Weights and Conclusion
doc.add_heading('Final Weights:', level=2)
doc.add_paragraph("W1 = 0.2")
doc.add_paragraph("W2 = 0.2")
doc.add_paragraph("Wb = -0.2")
doc.add_paragraph("")
doc.add_heading('Conclusion:', level=2)
doc.add_paragraph("After two iterations, the perceptron correctly classifies all input patterns. The perceptron converged with the final weights: W1 = 0.2, W2 = 0.2, and Wb = -0.2.")

# Save the document
file_path = "/mnt/data/Assignment_1_Solution.docx"
doc.save(file_path)
print("File saved at:", file_path)
